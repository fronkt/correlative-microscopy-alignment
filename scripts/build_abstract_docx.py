"""Render reports/abstract.{md,txt,docx} from the manuscript.

Title and abstract are PARSED FROM paper/paper.md rather than duplicated here.
An earlier version of this script hardcoded both, and they silently drifted: the
stored copy kept a retracted claim ("cross-modal appearance -- not scale -- is
what fails"), an intensifier the manuscript had already removed, and one-sided
p-values from before the two-sided correction. Parsing removes that failure mode.

Usage: python scripts/build_abstract_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PAPER = Path("paper/paper.md")

# Kept in step with the manuscript's "Hypothesis verdicts" subsection.
VERDICTS = [
    ("H1 (pyramid >= 35% gain at FOV <= 5%): ",
     "rejected; and untestable on this benchmark, which holds only four pairs "
     "below area ratio 0.05."),
    ("H2 (RoMa beats the ELoFTR family at low FOV): ", "supported."),
    ("H3 (affine sufficient vs. homography): ",
     "mostly supported (affine selected on 69% of well-registered pairs)."),
]

PROVENANCE = (
    "Full methods, tables, and figures: paper/paper.md. Metric-sensitivity and "
    "appearance-axis analyses: reports/metric_sensitivity.md and "
    "reports/appearance_axis.md. All numbers regenerate from results/ via the "
    "scripts referenced therein."
)


def strip_markdown(s: str) -> str:
    """Flatten inline markdown to plain text for the .txt/.docx renderings."""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    return s


def parse_paper() -> tuple[str, str]:
    if not PAPER.exists():
        sys.exit(f"{PAPER} not found; run from the repo root")
    text = PAPER.read_text(encoding="utf-8")

    m = re.search(r"^#\s+(.+?)\s*$", text, re.M)
    if not m:
        sys.exit("could not find the title (a leading '# ' line) in paper.md")
    title = m.group(1).strip()

    m = re.search(r"^##\s+Abstract\s*$(.+?)^---\s*$", text, re.M | re.S)
    if not m:
        sys.exit("could not find the '## Abstract' section in paper.md")
    body = "\n".join(
        ln.strip() for ln in m.group(1).strip().splitlines() if ln.strip()
    )
    if len(body) < 400:
        sys.exit(f"parsed abstract looks truncated ({len(body)} chars)")
    return title, body


def main() -> None:
    title, abstract_md = parse_paper()
    abstract_txt = strip_markdown(abstract_md)

    Path("reports").mkdir(exist_ok=True)

    md = [f"# {title}", "", "## Abstract", "", abstract_md, "",
          "## Hypothesis verdicts", ""]
    md += [f"- **{p.rstrip(': ')}**: {v}" for p, v in VERDICTS]
    md += ["", f"*{PROVENANCE}*", ""]
    Path("reports/abstract.md").write_text("\n".join(md), encoding="utf-8")

    txt = [title, "", "ABSTRACT", "", abstract_txt, "", "HYPOTHESIS VERDICTS", ""]
    txt += [f"  - {p}{v}" for p, v in VERDICTS]
    txt += ["", PROVENANCE, ""]
    Path("reports/abstract.txt").write_text("\n".join(txt), encoding="utf-8")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph().add_run("Abstract").bold = True
    body = doc.add_paragraph(abstract_txt)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph().add_run("Hypothesis verdicts").bold = True
    for prefix, verdict in VERDICTS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(prefix)
        p.add_run(verdict).bold = True

    doc.add_paragraph().add_run(PROVENANCE).italic = True
    doc.save("reports/abstract.docx")

    for f in ("reports/abstract.md", "reports/abstract.txt", "reports/abstract.docx"):
        print(f"wrote {f}")


if __name__ == "__main__":
    main()
